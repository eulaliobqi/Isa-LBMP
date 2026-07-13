# -*- coding: utf-8 -*-
"""Gera relatorio_isa_lbmp_didatico.docx a partir do conteudo final e revisado
em relatorio_isa_lbmp_didatico.md. Relatorio simplificado para leitor cientista
sem formacao em bioinformatica/filogenia (diferente do relatorio tecnico
completo gerado por gerar_relatorio_docx.py). Rodar de dentro do repo:
    python bin/gerar_relatorio_docx_didatico.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "results" / "results" / "results" / "figures"
OUT = ROOT / "relatorio_isa_lbmp_didatico.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=9)
        shade_cell(hdr[i], "1F3A5F")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=9)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def add_figure(doc, path, caption, width_cm=15):
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def code_block(doc, lines, size=8.5):
    """Bloco monoespacado para sequencias FASTA."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(size)
        if i < len(lines) - 1:
            run.add_break()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F1EC")
    pPr.append(shd)
    return p


doc = Document()

# ---- estilo base -----------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ---- capa ----------------------------------------------------------------
title = doc.add_heading(
    "Relatório simplificado — Busca por genes candidatos\nao florígeno (FT) em Urochloa", level=0
)
for run in title.runs:
    run.font.color.rgb = NAVY
p = doc.add_paragraph()
p.add_run(
    "Modulação da transição vegetativo-reprodutiva em Urochloa brizantha por RNA de "
    "interferência exógeno (LBMP/BIOAGRO-UFV, coordenação Drª Isabela Malaquias Dalto de Souza)"
).italic = True
para(doc, "")
p = doc.add_paragraph()
p.add_run("Este documento: ").bold = True
p.add_run(
    "versão simplificada e direta da etapa de bioinformática, escrita para um leitor "
    "cientista que não trabalha com bioinformática ou filogenia molecular. A versão "
    "técnica completa está em artigo.md e relatorio_isa_lbmp.docx."
)
p = doc.add_paragraph()
p.add_run("Data: ").bold = True
p.add_run("13 de julho de 2026")
doc.add_page_break()

# ---- 1. Apresentação ------------------------------------------------------
heading(doc, "1. Apresentação", level=1)
para(
    doc,
    "O florescimento em plantas é controlado, em parte, por uma proteína chamada FT "
    "(de FLOWERING LOCUS T), conhecida como \"florígeno\" — o sinal molecular que viaja "
    "da folha até o ápice da planta e dispara a troca da fase vegetativa para a "
    "reprodutiva. Em gramíneas forrageiras como Urochloa brizantha, atrasar essa troca "
    "prolonga a fase vegetativa (a fase de interesse para pastagem) e pode aumentar a "
    "produção de biomassa.",
)
para(
    doc,
    "O objetivo desta etapa do projeto foi encontrar, no material genético de Urochloa, "
    "qual gene corresponde ao FT — o alvo que depois será silenciado com RNA de "
    "interferência (dsRNA) aplicado por pulverização foliar. Sem saber exatamente qual "
    "sequência de DNA é o gene FT em Urochloa, não é possível desenhar o dsRNA.",
)
para(
    doc,
    "Este documento explica, em linguagem direta, como essa busca foi feita, o que foi "
    "encontrado e qual é o diagnóstico final: os candidatos mais fortes, com suas "
    "sequências completas.",
)

# ---- 2. Conceitos básicos -------------------------------------------------
heading(doc, "2. Conceitos básicos", level=1)
para(doc, "Antes da metodologia, quatro conceitos usados o tempo todo neste relatório:")
bullet(
    doc,
    "Gene ortólogo: dois genes de espécies diferentes são \"ortólogos\" quando descendem "
    "do mesmo gene ancestral e, em geral, mantêm a mesma função. Encontrar o ortólogo de "
    "FT em Urochloa significa achar o gene que, nessa espécie, faz o mesmo papel que o FT "
    "já bem conhecido em arroz, milho e Arabidopsis.",
)
bullet(
    doc,
    "Domínio proteico: um trecho da proteína com uma estrutura tridimensional "
    "característica, geralmente ligado a uma função específica. Duas proteínas que "
    "compartilham o mesmo domínio, na mesma posição da sequência, dificilmente chegaram "
    "a essa semelhança por acaso — é um indício forte de parentesco evolutivo (Seção 5).",
)
bullet(
    doc,
    "Genes \"parecidos, mas de função oposta\": a família à qual o FT pertence (chamada "
    "PEBP) inclui não só o FT (que promove o florescimento), mas também o TFL1 (que faz "
    "o oposto, mantém a planta vegetativa) e o MFT (envolvido em germinação de semente). "
    "Os três compartilham o mesmo domínio proteico — por isso não basta achar \"um gene "
    "parecido com FT\"; é preciso separar, dentro da família, qual ramo é FT de verdade "
    "e qual é TFL1/MFT (Seção 4.5).",
)
bullet(
    doc,
    "\"Candidato confirmado por bioinformática\" ≠ \"sequência pronta para uso em "
    "bancada\": tudo neste relatório vem de análise computacional (comparação de "
    "sequências, não de experimentos de laboratório). É uma evidência forte e "
    "reprodutível, mas a confirmação final exige PCR e sequenciamento em DNA real da "
    "planta (Seção 7).",
)

# ---- 3. Metodologia --------------------------------------------------------
heading(doc, "3. Metodologia, passo a passo", level=1)
para(
    doc,
    "A busca foi feita com um pipeline automatizado (Nextflow), executado sobre o "
    "genoma de Urochloa ruziziensis — montagem Embrapa_Uruz_1.0, acesso NCBI "
    "GCA_015476505.1, depositado pela Embrapa (nível cromossomo, isolado Clone 69, 9 "
    "cromossomos) — espécie diploide usada como substituta porque U. brizantha é "
    "poliploide e não tem genoma-referência próprio. Todo locus, coordenada e sequência "
    "apresentados neste relatório pertencem a esse genoma, salvo indicação explícita em "
    "contrário (Seção 6.1 traz a proveniência completa de cada sequência). As etapas, em ordem:",
)
numbered(
    doc,
    "Preparar 4 proteínas-modelo de FT já conhecidas — baixadas do banco público NCBI: "
    "Hd3a e RFT1 (arroz), ZCN8 (milho) e FT (Arabidopsis). São o \"gabarito\" usado para "
    "procurar o equivalente em Urochloa.",
)
numbered(
    doc,
    "Varredura ampla no genoma (TBLASTN) — comparação das 4 proteínas-modelo contra "
    "todo o genoma de U. ruziziensis, usando o programa BLAST no modo que compara "
    "proteína contra DNA. Limiar de rigor estatístico: e-value ≤ 1×10⁻⁵ (quanto menor o "
    "e-value, menor a chance de a semelhança encontrada ser coincidência).",
)
numbered(
    doc,
    "Confirmação cruzada por nucleotídeo (BLASTN) — mesma ideia, mas comparando a "
    "sequência de DNA codificante (CDS) das proteínas-modelo diretamente contra o "
    "genoma, com rigor maior: e-value ≤ 1×10⁻³.",
)
numbered(
    doc,
    "Reconstrução da estrutura do gene (miniprot) — programa que alinha proteína "
    "contra genoma respeitando a estrutura de éxons/íntrons dos genes de plantas "
    "(éxons = partes do gene mantidas na proteína final; íntrons = partes removidas "
    "antes da tradução — Seção 4.3), prevendo o \"recorte\" correto do gene e "
    "traduzindo a proteína resultante.",
)
numbered(
    doc,
    "Agrupamento em loci candidatos — os resultados das 3 etapas acima foram "
    "fundidos: hits a menos de 20 mil pares de base de distância foram tratados como "
    "pertencentes ao mesmo gene (\"locus\" candidato). Essa etapa gerou 24 loci "
    "candidatos espalhados pelos 9 cromossomos do genoma.",
)
numbered(
    doc,
    "Confirmação do domínio da família FT (HMMER) — cada uma das 24 proteínas "
    "candidatas foi comparada contra o perfil estatístico do domínio da família PEBP "
    "(código Pfam PF01161; Pfam é um banco público de perfis estatísticos de famílias "
    "de proteínas — Mistry et al., 2021), usando o programa HMMER (Eddy, 2011), com "
    "e-value ≤ 1×10⁻⁵. Só passaram os candidatos que realmente têm o domínio "
    "característico dessa família — 6 dos 24 loci confirmaram.",
)
numbered(
    doc,
    "Confirmação de ortologia real (RBH — \"reciprocal best hit\") — um candidato "
    "\"parecido\" com FT pode, na verdade, ser um parente próximo mas de função "
    "diferente (por exemplo, TFL1 ou MFT). Para reduzir esse risco, cada um dos 6 "
    "candidatos foi comparado contra o proteoma inteiro (o conjunto de todas as "
    "proteínas já descritas) de arroz, milho e Arabidopsis e vice-versa: só conta como "
    "confirmado quando o candidato aponta para o gene FT-modelo e o gene FT-modelo, "
    "ao procurar de volta em Urochloa, aponta para o mesmo candidato "
    "(\"reciprocidade\"). Rigor: e-value ≤ 1×10⁻¹⁰, mais rigoroso que a triagem "
    "inicial — a prática de aumentar a exigência estatística quando o objetivo passa "
    "de \"triagem ampla\" para \"confirmação 1-para-1\" segue a discussão de "
    "Moreno-Hagelsieb & Latimer (2008) sobre escolha de parâmetros de BLAST para "
    "detecção de ortólogos.",
)
numbered(
    doc,
    "Árvore genealógica (filogenia) — as sequências dos 6 candidatos confirmados "
    "foram alinhadas junto com as proteínas-modelo de FT, TFL1 e MFT usando o programa "
    "MAFFT (--auto; Katoh & Standley, 2013), o alinhamento foi aparado com trimAl "
    "(-automated1; Capella-Gutiérrez et al., 2009) para remover regiões pouco "
    "confiáveis, e a árvore foi construída com IQ-TREE 2 (-m MFP, 1000 réplicas de "
    "\"bootstrap\" -bb 1000 — bootstrap é uma medida de confiabilidade de 0 a 100 para "
    "cada agrupamento da árvore, explicada com mais detalhe na Seção 4.5; Minh et al., "
    "2020), usando as sequências de MFT como raiz da árvore — MFT é considerado o ramo "
    "mais antigo da família PEBP, presente até em musgos e licófitas, enquanto FT-like "
    "e TFL1-like não são (Karlgren et al., 2011). Esse passo é o que efetivamente "
    "separa \"é FT\" de \"é TFL1/MFT parecido\".",
)
numbered(
    doc,
    "Relatório final — consolidação de todos os resultados em tabelas e figuras (o "
    "que este documento apresenta).",
)
para(doc, "")
p = doc.add_paragraph()
p.add_run("Por que outra espécie (U. ruziziensis) e não U. brizantha diretamente? ").bold = True
p.add_run(
    "U. brizantha é uma espécie poliploide e de reprodução apomítica (agâmica), sem um "
    "genoma-referência montado cromossomo a cromossomo. U. ruziziensis é uma parente "
    "diploide próxima, com genoma bem montado, usada como substituta — prática padrão "
    "em genômica de espécies poliploides sem genoma próprio. Na Seção 4.6, esse ponto é "
    "revisitado: uma parte real do genoma de brizantha foi localizada depois, o que "
    "reduziu essa limitação."
)
add_figure(
    doc, FIG / "esquema_pipeline.png",
    "Figura 1. Fluxograma completo do pipeline, 9 etapas numeradas por função "
    "(descoberta de loci em azul, confirmação de ortologia em verde-água, filogenia e "
    "relatório em roxo).",
)

# ---- 4. Resultados ---------------------------------------------------------
heading(doc, "4. Resultados, etapa por etapa", level=1)

heading(doc, "4.1 Triagem ampla no genoma", level=2)
para(
    doc,
    "A varredura inicial (TBLASTN + BLASTN + miniprot) encontrou 24 loci candidatos, "
    "espalhados por 8 dos 9 cromossomos do genoma de U. ruziziensis.",
)
add_figure(
    doc, FIG / "loci_diagram.png",
    "Figura 2. Posição dos 24 loci candidatos ao longo dos cromossomos (eixo X = "
    "posição em pares de base). Vários candidatos aparecem próximos entre si no mesmo "
    "cromossomo (ex.: cromossomo CM027122.1, onde estão locus_009 e locus_012).",
)
add_table(
    doc, ["Item", "Valor"],
    [
        ["Total de loci candidatos (triagem ampla)", "24"],
        ["Cromossomos com pelo menos 1 candidato", "8 de 9"],
        ["Loci que confirmaram o domínio da família FT (etapa 6)", "6"],
        ["Loci que confirmaram ortologia por RBH (etapa 7)", "6 (os mesmos 6)"],
    ],
    col_widths_cm=[11, 5],
)
para(doc, "")

heading(doc, "4.2 Força da semelhança com os genes-modelo (bitscore)", level=2)
para(
    doc,
    "Cada locus recebeu uma pontuação de semelhança (\"bitscore\") contra os 4 "
    "genes-modelo — quanto maior, mais forte a semelhança.",
)
add_figure(
    doc, FIG / "bitscore_heatmap.png",
    "Figura 3. Mapa de calor do bitscore de cada um dos 24 loci contra os 4 genes de "
    "referência (colunas). locus_001 (linha do topo) se destaca claramente com o maior "
    "bitscore.",
)
add_table(
    doc, ["Locus", "Bitscore (TBLASTN)", "Proteína predita (aa)"],
    [
        ["locus_001", "155,0", "179"],
        ["locus_009", "129,0", "187"],
        ["locus_015", "129,0", "169"],
        ["locus_002", "123,0", "76 (fragmento)"],
        ["locus_005", "122,0", "175"],
        ["locus_012", "79,0", "72 (fragmento)"],
    ],
    col_widths_cm=[4, 6, 6],
)
para(doc, "")
para(
    doc,
    "Tamanhos acima são a predição bruta original do pipeline. A Seção 4.7 documenta "
    "uma correção encontrada depois: 3 dessas predições (locus_001, 005, 009) tinham "
    "alguns aminoácidos espúrios no início, já corrigidos nas sequências finais do "
    "Apêndice B.",
    italic=True, size=10,
)

heading(doc, "4.3 Estrutura do gene (éxons e íntrons)", level=2)
para(
    doc,
    "Genes de plantas são divididos em éxons (partes codificantes) separados por "
    "íntrons (partes removidas antes da tradução). A estrutura prevista para "
    "locus_001 tem 4 éxons e 3 íntrons — mesma contagem confirmada depois, de forma "
    "independente, no gene real equivalente encontrado no subgenoma real de brizantha "
    "(um \"subgenoma\" é o conjunto de cromossomos que uma espécie híbrida herdou de "
    "uma das duas espécies que a originaram — explicado em detalhe na Seção 4.6).",
)
add_figure(
    doc, FIG / "esquema_estrutura_genica.png",
    "Figura 4. Estrutura do gene em escala: predição em U. ruziziensis (só a região "
    "codificante, sem UTR) comparada com a anotação real do gene equivalente no "
    "subgenoma de brizantha — mesmos 4 éxons/3 íntrons, agora com as regiões "
    "não-codificantes (UTR) nas duas pontas.",
)

heading(doc, "4.4 Confirmação de que são genes reais, não só semelhantes (RBH)", level=2)
para(
    doc,
    "Um candidato \"parecido\" com FT pode ser, na verdade, um gene vizinho de função "
    "diferente. A confirmação por reciprocidade (etapa 7 da metodologia) checa isso em "
    "três espécies-modelo independentes:",
)
add_table(
    doc, ["Locus", "Confirma em arroz", "Confirma em milho", "Confirma em Arabidopsis", "Total"],
    [
        ["locus_001", "sim", "sim", "sim", "3/3"],
        ["locus_002", "sim", "sim", "sim", "3/3"],
        ["locus_005", "sim", "sim", "sim", "3/3"],
        ["locus_009", "sim", "sim", "sim", "3/3"],
        ["locus_012", "sim", "sim", "não", "2/3"],
        ["locus_015", "não", "não", "sim", "1/3"],
    ],
    col_widths_cm=[3, 3.5, 3.5, 4, 2],
)
para(doc, "")
para(
    doc,
    "locus_015 só confirma em Arabidopsis porque, em arroz e milho, o \"melhor hit de "
    "volta\" cai em locus_009 — um parálogo (cópia do mesmo gene, originada por "
    "duplicação dentro do próprio genoma de Urochloa) quase idêntico a locus_015. Esse "
    "tipo de resultado é uma limitação já documentada do método RBH especificamente em "
    "famílias gênicas com duplicação recente (Dalquen & Dessimoz, 2013): quando duas "
    "cópias são quase idênticas, o método tende a convergir sempre para a mesma cópia "
    "\"vencedora\".",
)
add_figure(
    doc, FIG / "esquema_convergencia_evidencias.png",
    "Figura 5. Síntese das linhas de evidência independentes que convergem para "
    "locus_001 — sequência (3 métodos: TBLASTN, BLASTN, miniprot), ortologia (RBH), "
    "filogenia (Seção 4.5) e confirmação no subgenoma real de brizantha (Seção 4.6).",
)

heading(doc, "4.5 Árvore genealógica: separando FT de TFL1/MFT", level=2)
para(
    doc,
    "A árvore (Seção 3, etapa 8) foi construída incluindo as proteínas-modelo de FT, "
    "TFL1 e MFT das 3 espécies-referência, com o grupo MFT como raiz.",
)
add_figure(
    doc, FIG / "phylogeny_tree_rooted_tfl1_mft.png",
    "Figura 6. Árvore genealógica enraizada. Os 6 loci confirmados se dividem em dois "
    "grupos distintos: 5 caem dentro do ramo FT (bootstrap 99 — quanto mais perto de "
    "100, mais confiável a separação) e 1 cai dentro do ramo MFT (bootstrap 100).",
)
add_table(
    doc, ["Locus", "Classificação final"],
    [
        ["locus_001", "FT-like (subgrupo Hd3a/RFT1)"],
        ["locus_002", "FT-like (evidência fraca, fragmento)"],
        ["locus_005", "FT-like (subgrupo ZCN8, o mais bem suportado — bootstrap 100)"],
        ["locus_009", "FT-like (par de cópias quase idênticas com locus_015)"],
        ["locus_012", "MFT-like — NÃO é o florígeno (agrupa com ZCN10 de milho, bootstrap 100; reclassificado)"],
        ["locus_015", "FT-like (par de cópias quase idênticas com locus_009)"],
    ],
    col_widths_cm=[3, 13],
)
para(doc, "")
para(
    doc,
    "Nenhum dos 6 loci caiu no ramo TFL1 (o antiflorígeno) — resultado tranquilizador, "
    "indicando que a triagem não trouxe \"antiflorígenos disfarçados\" entre os "
    "candidatos.",
)

heading(doc, "4.6 Confirmação em sequência real de U. brizantha", level=2)
para(
    doc,
    "U. brizantha não tem genoma-referência próprio, mas parte do genoma de uma "
    "espécie aparentada — U. decumbens cv. Basilisk (acesso NCBI/ENA GCA_964030465.3, "
    "BioProject PRJEB73762, Vega Group/Earlham Institute) — vem geneticamente de U. "
    "brizantha diploide. Isso acontece porque U. decumbens é uma espécie híbrida, "
    "originada do cruzamento de duas espécies diploides ancestrais: seu genoma é "
    "dividido em dois subgenomas — um conjunto de 18 cromossomos herdado de U. "
    "brizantha diploide, e outro conjunto de 18 cromossomos herdado de U. "
    "ruziziensis/U. decumbens diploide (confirmado no artigo que descreve esse "
    "genoma: de Jesus et al., bioRxiv 10.1101/2024.09.25.614935, PMC12005165). "
    "Comparando os 6 candidatos contra o proteoma desse genoma foi possível localizar "
    "o gene equivalente especificamente na porção que é subgenoma real de brizantha. "
    "(A tabela da Seção 6.1 resume, lado a lado, os dois genomas diferentes usados "
    "neste relatório — o de descoberta e o de confirmação — com seus números de "
    "acesso completos.)",
)
add_figure(
    doc, FIG / "esquema_subgenomas_decumbens.png",
    "Figura 7. Composição dos 36 cromossomos do genoma de U. decumbens cv. Basilisk "
    "por origem (18 vêm de brizantha diploide, 18 de ruziziensis/decumbens diploide), "
    "com destaque para as cópias do gene equivalente ao locus_001.",
)
add_table(
    doc, ["Locus", "Melhor cópia no subgenoma real de brizantha", "Identidade", "Gene ID"],
    [
        ["locus_001", "CAL5051155.1", "99,4% (1 aa de diferença em 174)", "URODEC1_LOCUS91967"],
        ["locus_002", "CAL5082316.1 / CAL5087382.1 (empatados)", "100,0%", "URODEC1_LOCUS109225 / URODEC1_LOCUS112169"],
        ["locus_005", "CAL4902238.1", "91,6% (maior divergência entre os 6)", "URODEC1_LOCUS9836"],
        ["locus_009", "CAL4952393.1", "98,8%", "URODEC1_LOCUS39484"],
        ["locus_012", "CAL4950975.1", "100,0%", "URODEC1_LOCUS38694"],
        ["locus_015", "CAL4952393.1 (mesma cópia de locus_009)", "98,8%", "URODEC1_LOCUS39484"],
    ],
    col_widths_cm=[3, 5.5, 4.5, 4],
)
para(doc, "")
para(
    doc,
    "Esse achado é importante: reduz a incerteza de \"espécie substituta\" para "
    "locus_001 a um número concreto e baixo (1 substituição de aminoácido), e mostrou "
    "que esse gene real tem estrutura de UTR (região não-codificante) anotada, algo "
    "que a predição inicial em U. ruziziensis não tinha.",
)

heading(doc, "4.7 Correção de um artefato de N-terminal nas predições (achado adicional)", level=2)
para(
    doc,
    "Comparar as predições contra os ortólogos reais (Seção 4.6) também revelou algo "
    "que passava despercebido: 3 das 6 proteínas preditas tinham alguns aminoácidos "
    "\"a mais\" no início, que não existem no gene real.",
)
p = doc.add_paragraph()
p.add_run("Como isso apareceu: ").bold = True
p.add_run(
    "a proteína predita de locus_001 não começa com metionina (M) — o aminoácido que "
    "sempre inicia uma proteína real (codificado pelo códon ATG). Comparando locus_001 "
    "com o gene real confirmado no subgenoma de brizantha (Seção 4.6), a sequência "
    "real bate exatamente com locus_001 a partir do 6º aminoácido — os 5 primeiros "
    "aminoácidos preditos não existem no gene de verdade."
)
p = doc.add_paragraph()
p.add_run("Causa: ").bold = True
p.add_run(
    "o programa miniprot alinha a proteína-modelo inteira (ex.: Hd3a, 179 posições) "
    "contra o genoma. Quando o início do gene real é um pouco diferente da "
    "proteína-modelo usada como referência, o programa às vezes estica o alinhamento "
    "alguns códons para dentro da região não-codificante (UTR) pra tentar completar o "
    "alinhamento inteiro — \"traduzindo\" um pedaço de UTR como se fosse parte da "
    "proteína."
)
para(
    doc,
    "Verificação nas outras 5 loci: baixei as proteínas reais dos ortólogos já "
    "identificados no subgenoma de brizantha para as outras 5 loci (mesma metodologia "
    "da Seção 4.6, arquivo brizantha_search/real_orthologs_decumbens.fasta) e comparei "
    "cada predição com seu ortólogo real:",
)
add_table(
    doc, ["Locus", "O que a comparação mostrou", "O que foi feito"],
    [
        ["locus_001", "5 aminoácidos a mais no início (LVAIG)", "Corrigido: proteína 179→174 aa, nucleotídeo 540→525 pb"],
        ["locus_005", "8 aminoácidos a mais no início (YNSTHPLV)", "Corrigido: proteína 175→167 aa, nucleotídeo 528→504 pb"],
        ["locus_009", "14 aminoácidos a mais no início (EDPSTRHRSASYRR)", "Corrigido: proteína 187→173 aa, nucleotídeo 564→522 pb"],
        ["locus_015", "Nenhum aminoácido a mais — faltam 4 aminoácidos reais no início", "Sem correção possível (não há como \"inventar\" a sequência que falta); mantido como está"],
        ["locus_002", "Fragmento corresponde à parte central de uma proteína real de 178 aa", "Sem modelo de gene — nada a corrigir, só documentar a correspondência"],
        ["locus_012", "Fragmento corresponde à parte central de uma proteína real de 171 aa", "Sem modelo de gene — nada a corrigir, só documentar a correspondência"],
    ],
    col_widths_cm=[2.5, 8, 6],
)
para(doc, "")
para(
    doc,
    "Cada correção foi verificada, não só aplicada: para as 3 loci corrigidas, "
    "confirmei que a sequência corrigida (1) começa com ATG, (2) traduz sem nenhum "
    "códon de parada no meio, (3) termina corretamente em um códon de parada, e (4) a "
    "proteína resultante bate exatamente com o esperado — se qualquer uma dessas "
    "checagens falhasse, a correção seria rejeitada automaticamente.",
)
para(
    doc,
    "O que isso muda na prática: os números \"179 aa\", \"175 aa\" e \"187 aa\" que "
    "aparecem nas tabelas de resultado do pipeline (Seção 4.2) continuam corretos "
    "como registro histórico do que o pipeline originalmente produziu. Mas as "
    "sequências corrigidas — 174, 167 e 173 aminoácidos — são as que devem ser "
    "usadas para qualquer desenho de primer ou dsRNA, por representarem o gene real "
    "com mais precisão. O Apêndice B já traz as sequências corrigidas.",
)

# ---- 5. Prova da homologia --------------------------------------------------
heading(doc, "5. Prova da homologia: o domínio compartilhado", level=1)
para(
    doc,
    "A evidência mais direta de que um gene pertence à família FT/TFL1/MFT é ele "
    "conter, na posição certa da proteína, o domínio característico dessa família "
    "(código Pfam PF01161, também chamado domínio PBP). Esse domínio tem 135 posições "
    "(\"aminoácidos\") no modelo estatístico de referência.",
)
para(
    doc,
    "A figura abaixo mostra, para cada um dos 6 candidatos, exatamente qual trecho da "
    "proteína alinhou com esse domínio (dado bruto da etapa 6 da metodologia, HMMER):",
)
add_figure(
    doc, FIG / "pebp_domain_coverage.png",
    "Figura 8. Cobertura do domínio PEBP (PF01161) por locus. A barra cinza no topo "
    "representa o domínio completo (135 aminoácidos). As barras coloridas mostram o "
    "trecho de cada candidato que efetivamente alinhou com esse domínio.",
    width_cm=16,
)
para(
    doc,
    "Leitura da figura: locus_001, 005, 009 e 015 (em azul) cobrem quase o domínio "
    "inteiro (das posições ~14–26 até ~100–111, dependendo do locus) — evidência "
    "ampla. locus_002 e 012 (em laranja) só cobrem a parte final do domínio (posições "
    "~79–135) — evidência mais restrita, consistente com serem fragmentos de proteína "
    "(76 e 72 aminoácidos, contra 169–187 dos outros quatro). Em nenhum dos 6 casos o "
    "alinhamento caiu fora da região do domínio — não é uma semelhança genérica de "
    "aminoácidos, é especificamente o trecho funcional da família FT/TFL1/MFT.",
)
para(
    doc,
    "Isso, combinado com a posição na árvore genealógica (Seção 4.5) e a confirmação "
    "de ortologia por RBH (Seção 4.4), é o que sustenta a afirmação de que esses 6 "
    "loci são realmente parentes evolutivos dos genes FT já conhecidos — não apenas "
    "proteínas com sequência parecida por acaso.",
)

# ---- 6. Diagnóstico final ----------------------------------------------------
heading(doc, "6. Diagnóstico final: as sequências candidatas", level=1)
para(doc, "Resumo compacto dos 6 loci confirmados. Sequências completas no Apêndice B.")

heading(doc, "6.1 De qual genoma vem cada sequência (rastreabilidade)", level=2)
para(
    doc,
    "Cada linha do diagnóstico combina dado de dois genomas diferentes, cada um com "
    "número de acesso público verificável — nenhum dos dois é um genoma de U. "
    "brizantha propriamente dito (ela não tem genoma-referência próprio, ver Seção 3):",
)
add_table(
    doc,
    ["", "Genoma", "Espécie", "Acesso", "Depositado por / banco", "Papel neste relatório"],
    [
        ["(A)", "Embrapa_Uruz_1.0", "U. ruziziensis (isolado Clone 69)", "GCA_015476505.1",
         "Embrapa — NCBI GenBank (nível cromossomo, 9 cromossomos)",
         "Genoma onde os 6 loci foram descobertos — origem das coordenadas CM027119.1–CM027126.1 e das sequências do Apêndice B"],
        ["(B)", "Assembly de U. decumbens cv. Basilisk",
         "U. decumbens (alotetraploide; 18 dos 36 cromossomos derivam de U. brizantha diploide)",
         "GCA_964030465.3",
         "Vega Group / Earlham Institute — ENA, BioProject PRJEB73762 (de Jesus et al., bioRxiv 10.1101/2024.09.25.614935, PMC12005165)",
         "Genoma onde os genes URODEC1_LOCUS... foram confirmados — fonte da coluna \"Gene ID\", restrita aos cromossomos com sufixo \"b\" (subgenoma real de brizantha)"],
    ],
    col_widths_cm=[1, 3, 4, 3, 4, 5],
)
para(doc, "")
para(
    doc,
    "Ou seja: a posição genômica sempre se refere ao genoma (A); o \"Gene ID\" sempre "
    "se refere a um gene real anotado no genoma (B), especificamente na porção desse "
    "genoma que descende de U. brizantha diploide.",
)
add_table(
    doc,
    ["Locus", "Gene ID — genoma (B), subgenoma brizantha", "Posição — genoma (A), U. ruziziensis", "Classificação", "Força da evidência"],
    [
        ["locus_001", "URODEC1_LOCUS91967 (99,4% idêntico)", "CM027126.1:30.720.251–30.721.504 (+)", "FT-like", "Muito forte — melhor candidato"],
        ["locus_002", "URODEC1_LOCUS109225 / 112169", "CM027119.1:57.431.465–57.436.138 (−)", "FT-like", "Fraca (fragmento, 76 aa)"],
        ["locus_005", "URODEC1_LOCUS9836", "CM027121.1:19.119.946–19.121.373 (+)", "FT-like", "Forte"],
        ["locus_009", "URODEC1_LOCUS39484", "CM027122.1:11.769.550–11.777.098 (+)", "FT-like", "Forte (par com locus_015)"],
        ["locus_012", "URODEC1_LOCUS38694", "CM027122.1:13.739.799–13.760.142 (+)", "MFT-like (não é florígeno)", "Moderada, mas classificação diferente"],
        ["locus_015", "URODEC1_LOCUS39484", "CM027123.1:10.494.327–10.498.283 (−)", "FT-like", "Moderada (par com locus_009; RBH só 1/3)"],
    ],
    col_widths_cm=[3, 4.5, 4.5, 3.5, 4.5],
)
para(doc, "")

heading(doc, "6.2 Descrição de cada candidato", level=2)
p = doc.add_paragraph()
p.add_run("locus_001 — o candidato principal. ").bold = True
p.add_run(
    "Reúne as 3 fontes de evidência de sequência simultaneamente, o maior bitscore "
    "(155,0), confirmação de ortologia nas 3 espécies-referência e posição robusta na "
    "árvore, dentro do subgrupo que contém Hd3a e RFT1 — os dois ortólogos de FT já "
    "caracterizados funcionalmente em arroz (Kojima et al., 2002). Foi confirmado, com "
    "99,4% de identidade, no subgenoma real de U. brizantha (Seção 4.6)."
)
p = doc.add_paragraph()
p.add_run("locus_002 — candidato fraco. ").bold = True
p.add_run(
    "Só 76 aminoácidos preditos (fragmento — não foi possível reconstruir a estrutura "
    "completa de éxons/íntrons), evidência de domínio parcial (Seção 5) e, no "
    "subgenoma real de brizantha, 3 cópias empatadas em identidade — não dá para "
    "apontar uma única cópia como \"a correta\" com o nível de evidência disponível."
)
p = doc.add_paragraph()
p.add_run("locus_005 — candidato forte, subgrupo diferente. ").bold = True
p.add_run(
    "Cai no subgrupo mais bem suportado da árvore (bootstrap 100), próximo ao ZCN8 de "
    "milho, gene da família FT já caracterizado por Danilevskaya et al. (2008). É o "
    "candidato com maior divergência da cópia real de brizantha (91,6% de identidade) "
    "— vale reavaliar se cogitado como alvo alternativo ao locus_001."
)
p = doc.add_paragraph()
p.add_run("locus_009 e locus_015 — par de cópias quase idênticas. ").bold = True
p.add_run(
    "Provavelmente resultado de uma duplicação gênica recente em Urochloa (o "
    "comprimento do ramo que separa os dois na árvore é quase zero). locus_009 "
    "confirma ortologia nas 3 espécies; locus_015, por ser tão parecido com locus_009 "
    "(Seção 4.4), só confirma em Arabidopsis. Os dois convergem para a mesma cópia "
    "real de brizantha (URODEC1_LOCUS39484)."
)
p = doc.add_paragraph()
p.add_run("locus_012 — reclassificado, não usar como alvo de FT. ").bold = True
p.add_run(
    "Inicialmente tratado como candidato fraco de FT, a árvore genealógica mostrou que "
    "na verdade agrupa (com suporte máximo, bootstrap 100) com o ZCN10 de milho — um "
    "gene MFT, não FT (Danilevskaya et al., 2008). MFT tem função relacionada à "
    "germinação de sementes, não ao florescimento. Mantido neste diagnóstico só para "
    "rastreabilidade — não deve ser considerado como alvo do gene FT-like da proposta."
)

# ---- 7. O que ainda falta ----------------------------------------------------
heading(doc, "7. O que ainda falta", level=1)
para(
    doc,
    "Um candidato confirmado por bioinformática não é, ainda, uma sequência pronta "
    "para desenho de dsRNA de bancada. Dois pontos seguem em aberto e são bloqueantes "
    "antes da síntese do dsRNA:",
)
numbered(
    doc,
    "Nenhuma validação experimental foi feita. Tudo neste relatório é análise "
    "computacional. O próximo passo recomendado é PCR seguido de sequenciamento "
    "(Sanger) em DNA genômico real da cultivar usada no projeto (cv. Marandu), usando "
    "como referência a sequência já confirmada no subgenoma real de brizantha "
    "(brizantha_search/brizantha_CDS.fasta).",
)
numbered(
    doc,
    "Especificidade (off-target) ainda não avaliada. A proposta do projeto exige "
    "checar se o dsRNA desenhado não vai silenciar, sem querer, outros genes "
    "parecidos — essa análise ainda não foi feita.",
)
para(doc, "Dois outros pontos, que antes eram bloqueantes, foram parcialmente resolvidos nesta rodada de análise:")
bullet(
    doc,
    "A diferença entre a espécie usada na busca (U. ruziziensis) e a espécie real do "
    "projeto (U. brizantha) era antes uma incerteza não quantificada — agora é um "
    "número concreto e baixo (1 substituição de aminoácido em 174, para o locus_001), "
    "graças à confirmação no subgenoma real (Seção 4.6). Ainda resta uma diferença "
    "possível entre a cultivar usada nessa confirmação (Basilisk) e a cultivar real do "
    "projeto (Marandu), não capturada aqui.",
)
bullet(
    doc,
    "A extensão exata do gene (incluindo as regiões não-codificantes, UTR) era "
    "desconhecida na predição original — agora há uma anotação real de UTR para o "
    "gene equivalente no subgenoma de brizantha (Seção 4.3), útil para desenhar "
    "primers fora da região codificante se necessário.",
)
para(
    doc,
    "Um ponto adicional, fora do escopo deste pipeline: o segundo gene-alvo do "
    "projeto, SBP5, ainda não tem sequência de referência identificada — a busca de "
    "loci equivalente para SBP5 não foi iniciada.",
)

# ---- 8. Referências -----------------------------------------------------------
heading(doc, "8. Referências", level=1)
refs = [
    "Capella-Gutiérrez, S., Silla-Martínez, J. M., & Gabaldón, T. (2009). trimAl: a tool for automated alignment trimming in large-scale phylogenetic analyses. Bioinformatics, 25(15), 1972–1973. https://doi.org/10.1093/bioinformatics/btp348",
    "Dalquen, D. A., & Dessimoz, C. (2013). Bidirectional Best Hits Miss Many Orthologs in Duplication-Rich Clades such as Plants and Animals. Genome Biology and Evolution, 5(10), 1800–1806. https://doi.org/10.1093/gbe/evt132",
    "Danilevskaya, O. N., Meng, X., Hou, Z., Ananiev, E. V., & Simmons, C. R. (2008). A Genomic and Expression Compendium of the Expanded PEBP Gene Family from Maize. Plant Physiology, 146, 250–264. https://doi.org/10.1104/pp.107.109538",
    "Eddy, S. R. (2011). Accelerated Profile HMM Searches. PLOS Computational Biology, 7(10), e1002195. https://doi.org/10.1371/journal.pcbi.1002195",
    "Karlgren, A., et al. (2011). Evolution of the PEBP Gene Family in Plants: Functional Diversification in Seed Plant Evolution. Plant Physiology, 156(4), 1967–1977. https://doi.org/10.1104/pp.111.176206",
    "Katoh, K., & Standley, D. M. (2013). MAFFT Multiple Sequence Alignment Software Version 7: Improvements in Performance and Usability. Molecular Biology and Evolution, 30(4), 772–780. https://doi.org/10.1093/molbev/mst010",
    "Kojima, S., Takahashi, Y., Kobayashi, Y., et al. (2002). Hd3a, a Rice Ortholog of the Arabidopsis FT Gene, Promotes Transition to Flowering Downstream of Hd1 under Short-Day Conditions. Plant and Cell Physiology, 43, 1096–1105. https://doi.org/10.1093/pcp/pcf156",
    "Minh, B. Q., Schmidt, H. A., Chernomor, O., Schrempf, D., Woodhams, M. D., von Haeseler, A., & Lanfear, R. (2020). IQ-TREE 2: New Models and Efficient Methods for Phylogenetic Inference in the Genomic Era. Molecular Biology and Evolution, 37(5), 1530–1534. https://doi.org/10.1093/molbev/msaa015",
    "Mistry, J., et al. (2021). Pfam: The Protein Families Database in 2021. Nucleic Acids Research, 49(D1), D412–D419. https://doi.org/10.1093/nar/gkaa913",
    "Moreno-Hagelsieb, G., & Latimer, K. (2008). Choosing BLAST options for better detection of orthologs as reciprocal best hits. Bioinformatics, 24(3), 319–324. https://doi.org/10.1093/bioinformatics/btm585",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(ref)
    run.font.size = Pt(10)

doc.add_page_break()

# ---- Apêndice A — Glossário ----------------------------------------------
heading(doc, "Apêndice A — Glossário", level=1)
bullet(doc, "BLAST / TBLASTN / BLASTN: família de programas que compara sequências (proteína ou DNA) para achar regiões semelhantes. TBLASTN compara proteína contra DNA; BLASTN compara DNA contra DNA.")
bullet(doc, "Bootstrap: em uma árvore genealógica, um número de 0 a 100 que indica a confiabilidade de um agrupamento — repete a análise centenas de vezes com pequenas variações nos dados e conta em quantas vezes o mesmo agrupamento aparece. Valores próximos de 100 indicam alta confiança.")
bullet(doc, "CDS: sigla de \"coding sequence\" — a parte de um gene que realmente é traduzida em proteína (sem os íntrons nem as regiões não-codificantes).")
bullet(doc, "Domínio proteico: ver Seção 2.")
bullet(doc, "E-value: medida estatística de quão provável é que uma semelhança encontrada seja coincidência. Quanto menor, mais confiável — 1×10⁻⁵ é bem mais rigoroso que 1×10⁻².")
bullet(doc, "Éxon / íntron: partes de um gene. Éxons são mantidos na proteína final; íntrons são removidos antes da tradução.")
bullet(doc, "Filogenia / árvore genealógica: representação gráfica de quão próximos evolutivamente estão diferentes genes ou espécies, com base em semelhança de sequência.")
bullet(doc, "HMM / perfil HMM: modelo estatístico que descreve, posição por posição, o padrão típico de uma família de proteínas — usado para reconhecer se uma nova sequência pertence a essa família.")
bullet(doc, "Ortólogo: ver Seção 2.")
bullet(doc, "Parálogo: cópia de um gene originada por duplicação dentro da mesma espécie (diferente de ortólogo, que vem de especiação).")
bullet(doc, "RBH (reciprocal best hit): método de confirmação de ortologia por reciprocidade — ver Seção 3, etapa 7.")
bullet(doc, "Subgenoma: em uma espécie poliploide originada do cruzamento de duas espécies diferentes, cada conjunto de cromossomos herdado de uma das duas espécies ancestrais.")
bullet(doc, "UTR: sigla de \"untranslated region\" — trechos de um gene transcritos em RNA, mas não traduzidos em proteína, localizados antes e depois da região codificante.")

doc.add_page_break()

# ---- Apêndice B — Sequências completas -------------------------------------
heading(doc, "Apêndice B — Sequências completas", level=1)
para(
    doc,
    "Todas as sequências abaixo (locus_001 a locus_015) foram extraídas do genoma (A) "
    "identificado na Seção 6.1: assembly Embrapa_Uruz_1.0 de Urochloa ruziziensis, "
    "acesso NCBI GCA_015476505.1, depositado pela Embrapa. As coordenadas de "
    "cromossomo em cada cabeçalho (CM027119.1–CM027126.1) são os próprios números de "
    "acesso GenBank desses cromossomos dentro desse assembly — não são identificadores "
    "internos do pipeline. Todas também estão disponíveis, em arquivo, em "
    "candidate_sequences/ (nucleotide_sequences.fasta e protein_sequences.fasta), "
    "junto com a metodologia de extração em candidate_sequences/README.md.",
)

heading(doc, "Sequências com CDS reconstruído, validado e corrigido (locus_001, 005, 009)", level=2)
para(
    doc,
    "Para essas 3 loci, a sequência de nucleotídeo foi reconstruída a partir do "
    "modelo gênico previsto (éxons concatenados na ordem genômica), validada, e "
    "corrigida removendo um artefato de N-terminal descoberto ao comparar cada "
    "predição contra o ortólogo real no subgenoma de brizantha (Seção 4.7): alguns "
    "aminoácidos que a predição bruta do miniprot colocava no início da proteína não "
    "existem no gene real — foram removidos, e a sequência abaixo já é a versão "
    "corrigida. A predição bruta original fica disponível em "
    "candidate_sequences/nucleotide_sequences.fasta e protein_sequences.fasta, com "
    "header _bruto_miniprot, para rastreabilidade.",
)

SEQS = [
    ("locus_001", "GCA_015476505.1", "CM027126.1:30.720.251–30.721.504 (+)",
     "corrigido: 525 pb", "corrigido: 174 aa", "540 pb / 179 aa", "5",
     ["ATGGCCAGGGACCCGTTGGTCGTTGGCAGGGTTGTGGGCGACGTGCTTGACCCCTTCGTCCGGACCACCA",
      "ACATGAGGGTCACCTTCGGCGGCAGGACCATCGCCAACGGCTGCGAGCTCAAGCCGTCCATGGTCACGCA",
      "CCAGCCCAGGGTCGATGTCGGCGGCCTCGACATGAGGACATTCTACACCCTTGTGATGGTCGACCCGGAT",
      "GCTCCAAGCCCAAGCGACCCCAACCTTAGGGAGTATTTGCACTGGCTGGTCACTGATATTCCAGGAACTA",
      "CCGGGGCAGCATTTGGGCAAGAGGTGATGTGCTACGAGAACCCTAGGCCGACCATGGGGATCCACCGCTT",
      "CGTGTTCGTGCTGTTCCAGCAGATGGGCCGGCAGACGGTGTACGCCCCCGGCTGGCGCCAGAACTTCAAC",
      "ACCAGGGACTTCGCCGAGCTCTACAACCTCGGCCCGCCGGTCGCCGCCGTCTACTTCAACTGCCAGCGCG",
      "AGGCAGGCTCCGGCGGCAGGAGGATGTACCCCTAA"],
     ["MARDPLVVGRVVGDVLDPFVRTTNMRVTFGGRTIANGCELKPSMVTHQPRVDVGGLDMRTFYTLVMVDPD",
      "APSPSDPNLREYLHWLVTDIPGTTGAAFGQEVMCYENPRPTMGIHRFVFVLFQQMGRQTVYAPGWRQNFN",
      "TRDFAELYNLGPPVAAVYFNCQREAGSGGRRMYP"]),
    ("locus_005", "GCA_015476505.1", "CM027121.1:19.119.967–19.121.373 (+)",
     "corrigido: 504 pb", "corrigido: 167 aa", "528 pb / 175 aa", "8",
     ["ATGGCTCATGGCATACAGGATGTGTTGGATCCCTTTACACCAACCGTTTCACTCAGAATAACATACAACA",
      "ACAGGCTACTTCTGTCAGGTGCTGAGCTAAAACCATCTGCTGTTGTAAGTAAACCGCGAGTTGATATTGG",
      "TGGAAATGACATGAGGGCTTTCTACACCCTGGTACTGATTGACCCAGATGCACCAAGTCCAAGCCATCCG",
      "TCACTAAGGGAGTACTTGCACTGGATGGTGACAGATATTCCTGAAACAACTAGCGCCAGCTTTGGCCAAG",
      "AGCTAGTGTTTTATGAGAAGCCAGAGCCAAGATCTGGTATCCACAGGATGGTATTTGTGCTGTTCCGGCA",
      "ACTTGGCAGGGGCACAGTTTTTGCACCAGAAATACGCCACCACTTCAACTGCAAAAGCTTTGCACAGCGA",
      "TATCACCTCAATATTGCCACCGCTACATATTTCAACTGTCAAAGGGAAGCTGGGTCGGGTGGAAGAAGGT",
      "TTAGAGATGAGTAG"],
     ["MAHGIQDVLDPFTPTVSLRITYNNRLLLSGAELKPSAVVSKPRVDIGGNDMRAFYTLVLIDPDAPSPSHP",
      "SLREYLHWMVTDIPETTSASFGQELVFYEKPEPRSGIHRMVFVLFRQLGRGTVFAPEIRHHFNCKSFAQR",
      "YHLNIATATYFNCQREAGSGGRRFRDE"]),
    ("locus_009", "GCA_015476505.1", "CM027122.1:11.769.550–11.777.098 (+)",
     "corrigido: 522 pb", "corrigido: 173 aa", "564 pb / 187 aa", "14",
     ["ATGCGGCGCGGCGACCCGCTGGTGGTGGGTCGGGTGATCGGCGACGTGGTGGACCCGTTCGTCCGGCGGG",
      "TGCCGCTGCGGGTGGCCTACACCGCGCGGGAGGTCTCCAACGGCTGCGAGCTCCGGCCATCCGCCGTCGC",
      "CGAGCAGCCGCGCGTCGAGGTCGGCGGGCCCGACATGAGAACCTTCTACACCCTGGTGCTGGTGGATCCG",
      "GACGCGCCAAGCCCCAGCAATCCCACCCTCAGGGAGTACTTGCACTGGCTGGTCACTGACATCCCGGCGA",
      "CGACAGGAGTTTCTTTTGGAACTGAGATCGTGTGCTACGAGAGCCCACGGCCGGTGCTCGGAATCCACCG",
      "GCTGGTGTTTCTGCTCTTCGAACAGCTCGGCCGGCAGACGGTCTACGCCCCAGGGTGGCGGCAGAACTTC",
      "AGCACCCGCGACTTCGCCGAGCTCTACAACCTCGGCCTGCCGGTCGCCGCCGTCTACTTCAATTGCCAAA",
      "GGGAGTCCGGAACTGGCGGGAGAAGAATGTGA"],
     ["MRRGDPLVVGRVIGDVVDPFVRRVPLRVAYTAREVSNGCELRPSAVAEQPRVEVGGPDMRTFYTLVLVDP",
      "DAPSPSNPTLREYLHWLVTDIPATTGVSFGTEIVCYESPRPVLGIHRLVFLLFEQLGRQTVYAPGWRQNF",
      "STRDFAELYNLGLPVAAVYFNCQRESGTGGRRM"]),
]

for locus_id, acc, coord, nt_size, aa_size, raw_size, n_removed, nt_lines, aa_lines in SEQS:
    p = doc.add_paragraph()
    p.add_run(f"{locus_id}").bold = True
    p.add_run(
        f" — genoma (A) {acc}, cromossomo {coord}, {nt_size}, proteína de {aa_size} "
        f"(predição bruta do miniprot: {raw_size}, com {n_removed} aminoácidos "
        f"espúrios no início — ver Seção 4.7)"
    )
    code_block(doc, [f">{locus_id} (nucleotídeo, CDS corrigido)"] + nt_lines + [""] + [f">{locus_id} (proteína, corrigida)"] + aa_lines)

heading(doc, "Sequência com N-terminal incompleto (locus_015)", level=2)
p = doc.add_paragraph()
p.add_run("locus_015").bold = True
p.add_run(
    " — genoma (A) GCA_015476505.1, cromossomo CM027123.1:10.494.327–10.498.271 (−), "
    "510 pb, proteína de 169 aa. Nota da verificação (Seção 4.7): comparado ao "
    "ortólogo real (CAL4952393.1, mesma cópia de locus_009), esta predição não tem "
    "nenhum aminoácido espúrio, mas está incompleta no início — falta o equivalente a "
    "4 aminoácidos reais (MRRG) que o miniprot não conseguiu alinhar. Sem correção "
    "possível; sequência mantida como veio do pipeline."
)
code_block(doc, [
    ">locus_015 (nucleotídeo, CDS reconstruído)",
    "GACCCGCTGGTGGTGGGTCGGGTGATCGGCGACGTGGTGGACCCGTTCGTCCGGCGGGTGCCGCTGCGG",
    "GTGGCCTACACCGCGCGGGAGGTCTCCAACGGCTGCGAGCTCCGGCCATCCGCCGTCGCCGAGCAGCCG",
    "CGCGTCGAGGTCGGCGGGCCCGACATGAGAACCTTCTACACCCTGGTGCTGGTAGATCCGGACGCGCCA",
    "AGCCCCAGCAATCCCACCCTCAGGGAGTACTTGCACTGGCTGGTCACTGACATCCCGGCGACGACAGGA",
    "GTTTCTTTTGGAACTGAGATCGTGTGCTACGAGAGCCCACGGCCGGTGCTCGGAATCCACCGGCTGGTG",
    "TTTCTGCTCTTCGAACAGCTCGGCCGGCAGACGGTCTACGCCCCAGGGTGGCGGCAGAACTTCAGCACC",
    "CGCGACTTCGCCGAGCTCTACAACCTCGGCCTGCCGGTCGCCGCCGTCTACTTCAATTGCCAAAGGGAG",
    "TCCGGAACTGGCGGGAGAAGAATGTGA",
    "",
    ">locus_015 (proteína)",
    "DPLVVGRVIGDVVDPFVRRVPLRVAYTAREVSNGCELRPSAVAEQPRVEVGGPDMRTFYTLVLVDPDAP",
    "SPSNPTLREYLHWLVTDIPATTGVSFGTEIVCYESPRPVLGIHRLVFLLFEQLGRQTVYAPGWRQNFST",
    "RDFAELYNLGLPVAAVYFNCQRESGTGGRRM",
])

heading(doc, "Sequências sem modelo de éxon (locus_002, locus_012)", level=2)
para(
    doc,
    "Importante: essas 2 loci têm apenas evidência de TBLASTN — o pipeline não "
    "conseguiu reconstruir um modelo de éxons/íntrons para elas (proteína incompleta, "
    "só um fragmento). Por isso, não existe uma sequência de CDS limpa disponível para "
    "essas duas. O que existe é o trecho genômico bruto completo entre as coordenadas "
    "de início e fim do locus — que quase certamente inclui íntron(s) e/ou regiões "
    "não-codificantes não removidas. Não colamos esses dois blocos aqui (4.674 pb e "
    "20.344 pb — inviável de ler em formato de texto), mas eles estão disponíveis, na "
    "íntegra, em candidate_sequences/nucleotide_sequences.fasta, com o aviso "
    "RAW_GENOMIC_SPAN exon_structure=UNKNOWN pode_incluir_introns_ou_UTR no cabeçalho.",
)
p = doc.add_paragraph()
p.add_run("locus_002").bold = True
p.add_run(
    " — genoma (A) GCA_015476505.1, cromossomo CM027119.1:57.431.465–57.436.138 (−), "
    "span bruto de 4.674 pb (proteína confirmada: 76 aa). Contexto da verificação "
    "(Seção 4.7): comparado ao ortólogo real (CAL5082316.1/CAL5087382.1, 178 aa), "
    "este fragmento corresponde às posições 101–173 dessa proteína — é a parte "
    "central/final do gene real, faltando tanto o início (100 aa) quanto o final "
    "(5 aa); os 3 primeiros aminoácidos do fragmento também não batem exatamente com "
    "o real."
)
code_block(doc, [
    ">locus_002 (proteína)",
    "TFSGNEIVPYESPRPPAGIHRIVFVLFKQQARQTVYAPGWRQNFNIRDFSAIYNLGAPVAALYFNCQKE",
    "SGVGGRR",
])
p = doc.add_paragraph()
p.add_run("locus_012").bold = True
p.add_run(
    " — genoma (A) GCA_015476505.1, cromossomo CM027122.1:13.739.799–13.760.142 (+), "
    "span bruto de 20.344 pb (proteína confirmada: 72 aa) — classificação: MFT-like, "
    "não é o florígeno (ver Seção 6). Contexto da verificação (Seção 4.7): comparado "
    "ao ortólogo real (CAL4950975.1, 171 aa), este fragmento corresponde exatamente "
    "às posições 100–171 dessa proteína (sem aminoácido espúrio), faltando os 99 "
    "aminoácidos iniciais do gene real."
)
code_block(doc, [
    ">locus_012 (proteína)",
    "GEEVVEYMGPRPPVGIHRYVLVLFEQKTRVNAEAPGERANFNTRAFAAAHELGLPTAVVYFNAQKEPAN",
    "RRR",
])

heading(doc, "Sequência real confirmada no subgenoma de U. brizantha (ortólogo de locus_001)", level=2)
para(
    doc,
    "Sequência real (não predita, diferente das 6 acima), do gene URODEC1_LOCUS91967 "
    "(accession de proteína CAL5051155.1), extraída do genoma (B) identificado na "
    "Seção 6.1: assembly de Urochloa decumbens cv. Basilisk, acesso GCA_964030465.3 "
    "(ENA, BioProject PRJEB73762, Vega Group/Earlham Institute; de Jesus et al., "
    "bioRxiv 10.1101/2024.09.25.614935, PMC12005165), especificamente no cromossomo "
    "36b (OZ075146.1:7.157.787–7.159.709) — a porção desse genoma confirmada, no "
    "artigo original do assembly, como derivada de U. brizantha diploide. 99,4% "
    "idêntica ao locus_001 (1 aminoácido de diferença em 174). CDS completo "
    "disponível em brizantha_search/brizantha_CDS.fasta; mRNA com UTR completo em "
    "brizantha_search/brizantha_mRNA_with_UTR.fasta; metodologia completa da busca "
    "(incluindo os outros 5 loci) em brizantha_search/README.md e "
    "brizantha_search/loci_brizantha_correlation.tsv.",
)

doc.save(OUT)
print(f"Salvo: {OUT}")
