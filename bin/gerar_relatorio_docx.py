# -*- coding: utf-8 -*-
"""Gera relatorio_isa_lbmp.docx a partir do conteudo consolidado e auditado
em artigo.md + resultados brutos do pipeline. Rodar de dentro do repo:
    python bin/gerar_relatorio_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "results" / "results" / "results" / "figures"
OUT = ROOT / "relatorio_isa_lbmp.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=9)
        shade_cell(hdr[i], "1F3A5F")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=9)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def add_figure(doc, path, caption, width_cm=15):
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


doc = Document()

# ---- estilo base -----------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ---- capa --------------------------------------------------------------
title = doc.add_heading(
    "Identificação e Validação In Silico de Ortólogos de\nFLOWERING LOCUS T (FT) em Urochloa spp.", level=0
)
for run in title.runs:
    run.font.color.rgb = NAVY
p = doc.add_paragraph()
p.add_run(
    "Relatório técnico — etapa de bioinformática do projeto \"Modulação da transição "
    "vegetativo-reprodutiva em Urochloa brizantha por RNA de interferência exógeno: "
    "validação funcional de SBP5 e FT-like e impacto na produção de biomassa\""
).italic = True
para(doc, "")
p = doc.add_paragraph()
p.add_run("Coordenação científica: ").bold = True
p.add_run("Drª Isabela Malaquias Dalto de Souza — Laboratório de Biologia Molecular de Plantas (LBMP), BIOAGRO/UFV")
p = doc.add_paragraph()
p.add_run("Análise de bioinformática: ").bold = True
p.add_run("pipeline Nextflow DSL2, repositório github.com/eulaliobqi/Isa-LBMP")
p = doc.add_paragraph()
p.add_run("Data do relatório: ").bold = True
p.add_run("11 de julho de 2026")
doc.add_page_break()

# ---- sumário executivo ---------------------------------------------------
heading(doc, "Sumário Executivo", level=1)
para(
    doc,
    "O pipeline de bioinformática identificou locus_001 (cromossomo CM027126.1, "
    "posição ~30.720.251–30.721.504, fita +, genoma de Urochloa ruziziensis) como "
    "candidato confirmado a ortólogo do gene FLOWERING LOCUS T (FT, florígeno), com três "
    "linhas de evidência independentes convergentes: (1) sequência — três fontes de "
    "evidência simultâneas (TBLASTN, BLASTN de CDS e miniprot), maior bitscore entre os "
    "candidatos (155,0); (2) ortologia — Reciprocal Best Hit (RBH) confirmado nas 3 "
    "espécies de referência (Oryza sativa, Zea mays, Arabidopsis thaliana), sem "
    "ambiguidade de parálogo; (3) filogenia — o locus cai dentro do clado FT-like "
    "(suporte estatístico bootstrap 99), separado do clado TFL1-like (bootstrap 98) e "
    "do agrupamento MFT-like (bootstrap 100), com essa separação interna robusta "
    "mesmo diante de uma limitação real de enraizamento identificada e documentada "
    "neste relatório (Seção 3.4).",
)
para(
    doc,
    "Um segundo achado relevante foi a reclassificação de locus_012: originalmente "
    "tratado como candidato FT-like fraco/fragmentado, a análise filogenética mostrou "
    "que na verdade se trata de um ortólogo de MOTHER OF FT AND TFL1 (MFT), agrupando "
    "com bootstrap 100 e comprimento de ramo próximo de zero à sequência ZCN10 de milho.",
)
para(
    doc,
    "Um TERCEIRO achado, obtido no mesmo dia da revisão que originou este relatório "
    "(11/07/2026), reduz diretamente a maior limitação apontada abaixo: não existe "
    "assembly próprio de U. brizantha, mas o assembly haplótipo-resolvido de "
    "Urochloa decumbens cv. Basilisk (GCA_964030465.3) tem cromossomos rotulados "
    "por origem de subgenoma, e 18 deles vêm de U. brizantha diploide real. Uma "
    "busca de locus_001 nesse proteoma encontrou uma cópia no subgenoma brizantha "
    "com 99,4% de identidade (1 aminoácido de diferença) e anotação de UTRs reais "
    "— informação que a predição original nunca teve (Seção 3.7).",
    bold=True,
)
para(
    doc,
    "IMPORTANTE — este relatório distingue explicitamente \"candidato confirmado por "
    "bioinformática\" de \"sequência pronta para desenho final de dsRNA em bancada\". "
    "A Seção 5 detalha os gargalos que devem ser resolvidos antes de qualquer síntese "
    "de dsRNA. Dois deles — divergência entre a espécie usada na análise original e a "
    "espécie-alvo real, e ausência de UTR — foram substancialmente reduzidos pelo "
    "achado acima. Seguem sem resolver: ausência de validação experimental na "
    "cultivar de trabalho (Marandu) e ausência da análise de especificidade "
    "off-target exigida pela proposta científica do projeto.",
)

# ---- 1. Introdução -------------------------------------------------------
heading(doc, "1. Introdução e Contexto", level=1)
para(
    doc,
    "O Brasil possui aproximadamente 160 milhões de hectares de pastagens cultivadas, "
    "com as gramíneas do gênero Urochloa (sin. Brachiaria) respondendo pela maior parte "
    "dessas áreas. Urochloa brizantha destaca-se como a principal forrageira do país, "
    "mas a produtividade das pastagens permanece abaixo do potencial devido à "
    "degradação de áreas de cultivo. O prolongamento da fase vegetativa — antes que "
    "recursos metabólicos sejam direcionados à reprodução — está associado ao aumento "
    "da produção de biomassa em diversas espécies cultivadas.",
)
para(
    doc,
    "O projeto de pesquisa que origina esta análise (proposta completa em "
    "projeto-isa.docx) visa avaliar o potencial da aplicação exógena de dsRNA livre "
    "(silenciamento transitório por pulverização foliar, não RNAi transgênico) "
    "direcionado aos genes SBP5 (família SPL, competência reprodutiva) e FT-like "
    "(florígeno) para modular a floração e aumentar a produção de biomassa em "
    "U. brizantha cv. Marandu.",
)
para(
    doc,
    "Este relatório cobre exclusivamente o Objetivo Específico 1 da proposta — "
    "\"Identificar regiões conservadas dos genes SBP5 e FT-like para o desenho de "
    "moléculas de dsRNA\" — restrito, por ora, ao gene FT-like. O gene SBP5 segue sem "
    "sequência de referência disponível (Seção 5.6).",
)
para(
    doc,
    "Como U. brizantha é uma espécie poliploide agâmica sem assembly cromossômico "
    "próprio, o genoma usado como alvo da busca é o de Urochloa ruziziensis (assembly "
    "Embrapa_Uruz_1.0, 9 cromossomos), espécie diploide filogeneticamente próxima, "
    "adotada como proxy. Esta escolha metodológica e suas limitações são discutidas em "
    "detalhe nas Seções 3.6 e 5.2.",
)

# ---- 2. Metodologia -------------------------------------------------------
heading(doc, "2. Metodologia", level=1)

heading(doc, "2.1 Pipeline de descoberta de loci", level=2)
para(
    doc,
    "Pipeline reprodutível em Nextflow DSL2 (11 processos, repositório "
    "github.com/eulaliobqi/Isa-LBMP), executado em servidor Debian (32 núcleos), com "
    "ambientes conda/mamba criados automaticamente por processo:",
)
add_table(
    doc,
    ["#", "Processo", "Função", "Parâmetro-chave"],
    [
        ["1", "BUILD_BLASTDB", "Indexa o genoma alvo (makeblastdb)", "—"],
        ["2", "TBLASTN", "Triagem ampla proteína × genoma", "E-value 1e-5"],
        ["2b", "BLASTN_CDS", "Evidência cruzada CDS nucleotídeo × genoma", "E-value 1e-3"],
        ["3", "MINIPROT_INDEX/ALIGN", "Estrutura gênica spliced-aware + tradução", "—"],
        ["4", "CONSOLIDATE_LOCI", "Funde hits das 3 fontes em loci candidatos", "cluster <20 kb"],
        ["5", "HMMSEARCH_PEBP", "Confirma domínio PEBP (PF01161) via HMMER", "E-value 1e-5"],
        ["6", "DOWNLOAD_REF_PROTEOMES", "Baixa proteomas de arroz/milho/Arabidopsis", "assemblies RefSeq"],
        ["7", "RECIPROCAL_BEST_HIT", "BLASTP recíproco — confirma ortologia", "E-value 1e-10"],
        ["8", "PHYLOGENY", "MAFFT + trimAl + IQ-TREE", "ver 2.3"],
        ["9", "REPORT", "Tabela final + figuras", "—"],
    ],
    col_widths_cm=[1.2, 4, 7, 3],
)
para(doc, "")
para(
    doc,
    "Sequências de query FT-like (reference_queries.fasta): 4 proteínas de referência "
    "do NCBI — Hd3a (NP_001408118.1) e RFT1 (NP_001408117.1) de Oryza sativa, ZCN8 "
    "(NP_001106247.1) de Zea mays, FT (NP_001320342.1) de Arabidopsis thaliana.",
)

heading(doc, "2.2 Expansão da referência com homólogos reais de Urochloa", level=2)
para(
    doc,
    "Para dar contexto filogenético mais próximo ao alvo, rodou-se BLASTP remoto (NCBI, "
    "banco nr) para cada uma das 4 proteínas de referência, restrito a Urochloa[Organism], "
    "baixando as 5 sequências distintas mais próximas de cada uma. Resultado: 20 hits, "
    "reduzidos a 15 sequências únicas após remoção de duplicatas — todas de U. humidicola "
    "e U. decumbens (72,8–88,4% de identidade); nenhuma de U. ruziziensis/U. brizantha "
    "especificamente, refletindo a escassez de proteínas anotadas dessas duas espécies "
    "em bancos públicos.",
)

heading(doc, "2.3 Filogenia com discriminação FT vs. TFL1/MFT", level=2)
para(
    doc,
    "O domínio PEBP (PF01161) é compartilhado por FT (florígeno), TFL1 (antiflorígeno) "
    "e MFT — genes de função oposta. Uma filogenia construída só com referências FT-like "
    "não consegue, por si, confirmar que um candidato é FT e não um parálogo antagonista. "
    "Foram então identificadas, por BLASTP verificado (E-value 1e-10, restrito a "
    "Oryza sativa/Zea mays), sequências TFL1-like e MFT-like reais, usando como query "
    "TFL1 (NP_196004.1) e MFT (NP_173250.1) de Arabidopsis — ambos confirmados pelo "
    "campo [Gene Name] do NCBI, não por busca textual livre (que colidiu: o símbolo "
    "histórico RCN1 de arroz foi reatribuído no NCBI atual a um gene não relacionado, "
    "um transportador ABC).",
)
add_table(
    doc,
    ["Família", "Arabidopsis", "Arroz", "Milho"],
    [
        ["TFL1-like", "NP_196004.1", "XP_015624118.1 (SELF-PRUNING/CEN-like, 72,8% id.)", "NP_001106241.1 (ZCN2, 74,0% id.)"],
        ["MFT-like", "NP_173250.1", "NP_001410922.1 (66,9% id.)", "NP_001106249.1 (ZCN10, 63,0% id.)"],
    ],
    col_widths_cm=[2.5, 3, 6, 4],
)
para(doc, "")
para(
    doc,
    "Conjunto de referência final para a filogenia (reference_queries_expanded_tfl1_mft.fasta, "
    "25 sequências) = 4 FT-like originais + 15 homólogos de Urochloa + 3 TFL1-like + 3 MFT-like. "
    "Alinhamento MAFFT (--auto) → aparo trimAl (-automated1) → IQ-TREE (-m MFP, -bb 1000, -T AUTO), "
    "com as 3 sequências MFT-like especificadas como outgroup (ver Seção 3.4 para a avaliação "
    "crítica do sucesso desse enraizamento).",
)

heading(doc, "2.4 Justificativa metodológica (síntese)", level=2)
para(
    doc,
    "Cada escolha metodológica relevante foi checada contra a literatura, com verificação "
    "da fonte real (WebSearch + WebFetch) antes de qualquer citação entrar neste relatório "
    "— nenhuma referência foi usada de memória sem confirmação. A lista completa de 19 "
    "citações com a justificativa detalhada de cada escolha (thresholds de E-value, "
    "ferramentas, e uma explicação da literatura para por que o método Reciprocal Best Hit "
    "pode falhar diante de parálogos quase idênticos — exatamente o caso observado entre "
    "locus_009 e locus_015) está disponível no documento de trabalho artigo.md, Seção 6, "
    "e resumida na Seção 6 (Referências) deste relatório. Dois destaques:",
)
bullet(
    doc,
    "Dalquen & Dessimoz (2013, Genome Biology and Evolution) demonstram que o método RBH "
    "pode perder até ~60% das relações de ortologia verdadeiras em clados ricos em "
    "duplicação gênica — evidência direta de por que locus_015 (parálogo quase idêntico "
    "a locus_009) só confirma RBH em 1 de 3 espécies: limitação conhecida do método, não "
    "falha do pipeline.",
)
bullet(
    doc,
    "Karlgren et al. (2011, Plant Physiology) e de Jesus et al. (2022, Frontiers in "
    "Genetics) sustentam a lógica de MFT como grupo basal ao split FT/TFL1 na família "
    "PEBP — a base biológica para tentar usar MFT como outgroup nesta análise (ver "
    "Seção 3.4 para os limites dessa tentativa nos dados deste projeto).",
)

# ---- 3. Resultados ---------------------------------------------------------
heading(doc, "3. Resultados", level=1)

heading(doc, "3.1 Loci candidatos e confirmação por domínio PEBP", level=2)
para(
    doc,
    "24 loci candidatos foram identificados na triagem ampla (TBLASTN). Destes, 6 "
    "confirmam o domínio PEBP via HMMER: locus_001, 002, 005, 009, 012 e 015.",
)
add_figure(
    doc,
    FIG / "loci_diagram.png",
    "Figura 1. Distribuição dos 24 loci candidatos ao longo dos 9 cromossomos do "
    "assembly Embrapa_Uruz_1.0 de Urochloa ruziziensis, por posição genômica (pb).",
    width_cm=15,
)

heading(doc, "3.2 Reciprocal Best Hit (RBH) — confirmação de ortologia", level=2)
para(
    doc,
    "Após correção de um bug de normalização de identificador (commit 9f96ac6), uma "
    "rodagem completa sem cache confirmou os seguintes resultados:",
)
add_table(
    doc,
    ["Locus", "Evidências", "Bitscore", "Proteína", "PEBP", "RBH (3 espécies)"],
    [
        ["locus_001", "tblastn+miniprot+blastn_cds", "155,0", "179 aa", "Sim", "3/3 confirmado"],
        ["locus_009", "tblastn+miniprot", "129,0", "187 aa", "Sim", "3/3 confirmado"],
        ["locus_015", "tblastn+miniprot", "129,0", "169 aa", "Sim", "1/3 (só Arabidopsis)"],
        ["locus_002", "tblastn", "123,0", "76 aa (fragmento)", "Sim", "confirmado (evidência fraca)"],
        ["locus_005", "tblastn+miniprot", "122,0", "175 aa", "Sim", "3/3 confirmado"],
        ["locus_012", "tblastn", "79,0", "72 aa (fragmento)", "Sim", "2/3"],
    ],
    col_widths_cm=[2.2, 4, 2, 3.3, 1.5, 4],
)
para(doc, "")
para(
    doc,
    "locus_015 só confirma reciprocidade com Arabidopsis — em arroz e milho, o melhor "
    "hit reverso cai em locus_009 (parálogo quase idêntico, comprimento de ramo próximo "
    "de zero na árvore filogenética), sinal biológico de duplicação em tandem recente, "
    "consistente com a literatura sobre limitações do método RBH (Seção 2.4).",
)
add_figure(
    doc,
    FIG / "bitscore_heatmap.png",
    "Figura 2. Mapa de calor do bitscore (TBLASTN) de cada um dos 24 loci candidatos "
    "contra as 4 proteínas de referência convergentes. Valores mais altos (amarelo) "
    "indicam maior similaridade de sequência.",
    width_cm=9,
)

heading(doc, "3.3 Filogenia — três clados FT-like distintos", level=2)
para(
    doc,
    "A análise filogenética revelou três clados FT-like distintos em Urochloa, um "
    "padrão biologicamente esperado — gramíneas costumam ter pequenas famílias de "
    "genes FT-like (arroz: Hd3a+RFT1; milho: múltiplos genes ZCN):",
)
bullet(
    doc,
    "locus_001 + homólogos de Urochloa recuperados via BLASTP contra Hd3a/RFT1 + o "
    "par Hd3a/RFT1 de arroz (bootstrap 82) — candidato mais próximo do florígeno "
    "clássico tipo Hd3a/RFT1.",
)
bullet(
    doc,
    "locus_009/locus_015 (bootstrap 100 entre si, quase idênticos) + homólogos "
    "recuperados via BLASTP contra Arabidopsis-FT, próximo (bootstrap 56) de "
    "NP_001320342.1 e dos fragmentos curtos locus_002/012.",
)
bullet(
    doc,
    "locus_005 + NP_001106247.1 (ZCN8) + homólogos recuperados via BLASTP contra "
    "ZCN8 — bootstrap 100, o clado mais bem suportado do conjunto.",
)

heading(doc, "3.4 Filogenia com outgroup MFT-like — resultado e ressalva crítica", level=2)
para(
    doc,
    "Uma rodagem dedicada incluiu as 3 sequências MFT-like como outgroup do IQ-TREE, "
    "na tentativa de enraizar a árvore com base biológica (MFT é considerado grupo "
    "basal ao split evolutivo FT/TFL1). O resultado confirma a separação entre os "
    "três grupos esperados, mas com uma ressalva metodológica importante identificada "
    "em auditoria dedicada:",
)
para(
    doc,
    "ACHADO DE AUDITORIA: o log do IQ-TREE (urochloa_ft_phylo.iqtree) declara "
    "textualmente: \"NOTE: Tree is UNROOTED although outgroup taxon 'NP_173250.1' is "
    "drawn at root\". A raiz da árvore é, de fato, uma tricotomia não resolvida de 3 "
    "ramos, e as 3 sequências MFT-like especificadas como outgroup não formam um clado "
    "monofilético nesta topologia: NP_173250.1 (MFT de Arabidopsis) ocupa um ramo "
    "terminal isolado, separado do par NP_001410922.1 (arroz) + NP_001106249.1 (milho).",
    italic=True,
)
para(
    doc,
    "O que isso muda, e o que não muda: a conclusão central deste relatório não "
    "depende de resolver essa tricotomia basal. Os três ramos da raiz correspondem a "
    "(a) o clado FT-like+TFL1-like combinado — 27 sequências, suporte estatístico "
    "bootstrap 100, bem resolvido e monofilético; (b) o par locus_012+NP_001106249.1 "
    "(bootstrap 100) associado a NP_001410922.1 (bootstrap 62); e (c) NP_173250.1 "
    "isolada. Dentro do ramo (a), a separação entre o clado FT-like (bootstrap 99) e "
    "o clado TFL1-like (bootstrap 98) é robusta e independente de qual dos três ramos "
    "da raiz é tratado como ancestral. O que não está comprovado é a alegação mais "
    "forte de enraizamento biológico completo e limpo por MFT — essa está apenas "
    "parcialmente resolvida.",
)
add_figure(
    doc,
    FIG / "phylogeny_tree_rooted_tfl1_mft.png",
    "Figura 3. Árvore de máxima verossimilhança (IQ-TREE, modelo MFP, 1000 réplicas "
    "de bootstrap ultrarrápido) das 27 sequências FT-like/TFL1-like mais as 3 "
    "sequências MFT-like usadas como outgroup pretendido. Valores nos nós internos "
    "são suportes de bootstrap (%). Nota: a raiz é uma tricotomia não resolvida — "
    "ver ressalva no texto (Seção 3.4).",
    width_cm=16,
)
para(
    doc,
    "locus_001 cai dentro do clado FT-like (bootstrap 99), separado do TFL1-like "
    "(bootstrap 98) e do agrupamento MFT-like que contém locus_012 (bootstrap 100 "
    "para esse agrupamento) — confirmado pela topologia interna, independentemente "
    "da tricotomia basal não resolvida.",
    bold=True,
)

heading(doc, "3.5 Reclassificação de locus_012", level=2)
para(
    doc,
    "locus_012, originalmente tratado como candidato FT-like fraco (evidência apenas "
    "TBLASTN, proteína curta de 72 aa), na verdade não é FT-like: agrupa com "
    "bootstrap 100 e comprimento de ramo próximo de zero à sequência ZCN10 de milho "
    "(NP_001106249.1, confirmada MFT-like por BLASTP verificado), sister à sequência "
    "MFT de arroz (NP_001410922.1). Esse achado é robusto independentemente da "
    "ressalva sobre a tricotomia basal (Seção 3.4), pois o agrupamento locus_012 + "
    "ZCN10 está inteiramente dentro do ramo (b) da raiz, não afetado pela posição de "
    "NP_173250.1.",
)

heading(doc, "3.6 Espécie-proxy: Urochloa ruziziensis", level=2)
para(
    doc,
    "Todos os resultados das Seções 3.1 a 3.5 referem-se a loci identificados no "
    "genoma de Urochloa ruziziensis, usado como proxy diploide na ausência de "
    "assembly próprio de U. brizantha. Nenhuma das 15 sequências homólogas "
    "recuperadas por BLASTP (Seção 2.2) pertence a U. ruziziensis ou U. brizantha "
    "especificamente — todas vêm de U. humidicola e U. decumbens, refletindo a "
    "escassez de proteínas anotadas das duas primeiras espécies em bancos públicos. "
    "Esta limitação foi substancialmente mitigada pelo achado da Seção 3.7.",
)

heading(doc, "3.7 Confirmação em subgenoma real de U. brizantha (achado de 11/07/2026)", level=2)
para(
    doc,
    "Não existe assembly próprio de U. brizantha, mas existe algo melhor que o "
    "proxy usado nas Seções 3.1–3.6: o assembly haplótipo-resolvido, "
    "cromossomo-completo de Urochloa decumbens cv. Basilisk (GCA_964030465.3, de "
    "Vega Group/Earlham Institute, ENA, BioProject PRJEB73762) é um alotetraploide "
    "cujos 36 cromossomos foram rotulados por origem de subgenoma, confirmado no "
    "próprio texto do artigo original: sufixo \"b\" = subgenoma derivado de "
    "U. brizantha diploide; sufixo \"rd\" = subgenoma derivado de "
    "U. ruziziensis/U. decumbens diploide.",
)
para(
    doc,
    "BLASTP local (BLAST+ 2.17.0) da proteína de locus_001 contra o proteoma "
    "anotado de U. decumbens (150.710 proteínas) encontrou 5 cópias quase "
    "idênticas (99,4–100% identidade), refletindo a redundância de haplótipos do "
    "assembly:",
)
add_table(
    doc,
    ["Accession", "Identidade", "Cromossomo", "Subgenoma"],
    [
        ["CAL5036879.1", "100,0%", "33rd", "ruziziensis/decumbens"],
        ["CAM0151951.1", "99,4%", "scaffold não localizado", "—"],
        ["CAM0146821.1", "99,4%", "scaffold não localizado", "—"],
        ["CAL5051155.1", "99,4%", "36b", "brizantha (real)"],
        ["CAL5044486.1", "99,4%", "34rd", "ruziziensis/decumbens"],
    ],
    col_widths_cm=[3, 2.5, 5, 5],
)
para(doc, "")
para(
    doc,
    "CAL5051155.1 (gene URODEC1_LOCUS91967, cromossomo 36b, "
    "OZ075146.1:7.157.787–7.159.709, fita +) é a cópia real do subgenoma "
    "brizantha e difere da predição de locus_001 (via U. ruziziensis) por "
    "exatamente 1 aminoácido (posição ~53, V→I) — 173/174 aa idênticos.",
    bold=True,
)
para(
    doc,
    "Dois ganhos diretos: (1) a divergência espécie-proxy, antes desconhecida, "
    "agora é um número baixo e conhecido — 1 substituição conservativa em 174 "
    "posições; (2) o gene tem anotação de mRNA completo com UTRs reais (1.222 nt "
    "totais vs. 525 nt de CDS, ~697 nt de UTR 5'+3'), algo que a predição via "
    "miniprot em U. ruziziensis nunca teve. A estrutura de 4 éxons/3 íntrons foi "
    "confirmada de forma independente, reforçando a confiança na predição "
    "original.",
)
para(
    doc,
    "O que isso não resolve: a cv. Basilisk (U. decumbens) não é a cv. Marandu "
    "(U. brizantha, a cultivar real do projeto) — pode haver diferença adicional "
    "entre cultivares/acessos não capturada aqui. Confirmação por PCR e "
    "sequenciamento Sanger em DNA genômico real da cv. Marandu continua "
    "recomendada (Seção 5.1), mas agora partindo de uma sequência-molde de "
    "U. brizantha de verdade, com risco de divergência residual já quantificado "
    "como baixo. Sequências completas disponíveis em brizantha_search/ no "
    "repositório do projeto.",
)

# ---- 4. Discussão -----------------------------------------------------------
heading(doc, "4. Discussão", level=1)
bullet(
    doc,
    "locus_001 é o candidato confirmado a ortólogo de FT em U. ruziziensis dentro do "
    "que a análise de bioinformática pode provar: única confirmação com as três "
    "fontes de evidência simultâneas, maior bitscore, predição de proteína de 179 aa "
    "sem stop códon prematuro (mas sem UTRs modeladas e sem confirmação por RNA-seq "
    "real de Urochloa — a predição vem inteiramente de alinhamento spliced-aware de "
    "proteínas de outras espécies via miniprot), RBH limpo nas 3 espécies e posição "
    "robusta na filogenia. É um candidato sólido para reportar como resultado de "
    "bioinformática — não é, ainda, uma sequência pronta para desenho final de "
    "dsRNA/primers de bancada sem validação experimental adicional (Seção 5).",
)
bullet(
    doc,
    "A existência de três clados FT-like distintos dentro do ingroup confirmado é "
    "biologicamente esperada e reforça que a triagem ampla capturou a diversidade "
    "real do locus gênico, não ruído estatístico.",
)
bullet(
    doc,
    "A reclassificação de locus_012 (de \"FT-like fraco\" para MFT-like verdadeiro) "
    "foi confirmada pela topologia da árvore, independentemente da ressalva sobre "
    "enraizamento — o gene MFT regula germinação e dormência de semente e integra "
    "sinalização floral de forma indireta; não é o florígeno em si e não deve ser "
    "usado como alvo do gene FT-like da proposta.",
)
bullet(
    doc,
    "locus_002 (76 aa, fragmento curto, só evidência TBLASTN) segue com evidência "
    "mais fraca dentro do clado FT-like — pode ser um gene real truncado na "
    "montagem, um pseudogene, ou artefato de triagem; não recomendado como candidato "
    "primário.",
)

# ---- 5. Limitações e gargalos --------------------------------------------
heading(doc, "5. Limitações e Gargalos para Aplicação em Bancada", level=1)
para(
    doc,
    "Esta seção existe porque candidato bioinformático confirmado não é o mesmo que "
    "sequência pronta para síntese de dsRNA. Revisão crítica dedicada identificou os "
    "gargalos abaixo, em ordem de prioridade. Os itens 5.1 a 5.3 são bloqueantes: "
    "recomenda-se não avançar para desenho final de primers/dsRNA sem resolvê-los.",
)

heading(doc, "5.1 Zero validação experimental até agora (bloqueante)", level=2)
para(
    doc,
    "Todo o resultado apresentado é in silico. Nenhuma base do locus foi confirmada "
    "por PCR em DNA genômico real; nenhuma expressão foi confirmada por RT-PCR/RT-qPCR.",
)
para(
    doc,
    "Recomendação: antes de qualquer síntese de dsRNA, amplificar por PCR (par de "
    "primers tolerante, ~200–400 pb dentro do CDS predito) a região homóloga a "
    "locus_001 em DNA genômico real de U. brizantha cv. Marandu e sequenciar por "
    "Sanger. Realizar também RT-PCR de folha para confirmar transcrição e que a "
    "junção éxon-éxon prevista corresponde ao cDNA real. Definir a sequência final "
    "do dsRNA somente após essa etapa.",
    bold=True,
)

heading(doc, "5.2 Gap espécie-proxy — PARCIALMENTE RESOLVIDO em 11/07/2026 (era bloqueante)", level=2)
para(
    doc,
    "locus_001 foi originalmente encontrado só em U. ruziziensis (diploide), não "
    "em U. brizantha (poliploide agâmica, a espécie real de bancada), sem nenhuma "
    "estimativa de divergência entre as duas.",
)
para(
    doc,
    "Atualização (Seção 3.7): o subgenoma brizantha real do assembly de "
    "U. decumbens cv. Basilisk (GCA_964030465.3) tem uma cópia quase idêntica "
    "(CAL5051155.1, cromossomo 36b) — 99,4% idêntica, 1 substituição de "
    "aminoácido em 174. A divergência deixou de ser uma incógnita e passou a ser "
    "um número baixo e conhecido.",
    bold=True,
)
para(
    doc,
    "O que ainda falta: essa é a cultivar Basilisk, não a Marandu (cultivar real "
    "do projeto) — pode haver diferença adicional entre cultivares/acessos de "
    "brizantha não capturada aqui. Recomendação atualizada: usar a sequência "
    "brizantha_CDS.fasta (Seção 3.7) — não mais a predição em U. ruziziensis "
    "isolada — como molde para desenho de primers de triagem; confirmar por PCR e "
    "Sanger em DNA genômico real da cv. Marandu (Seção 5.1) antes do desenho "
    "final, agora partindo de uma base de brizantha de verdade com risco de "
    "divergência residual já conhecido como baixo.",
    bold=True,
)

heading(doc, "5.3 Off-target ainda não avaliado — requisito da própria proposta (bloqueante)", level=2)
para(
    doc,
    "A proposta científica do projeto exige explicitamente análise de "
    "especificidade do dsRNA contra o transcriptoma/genoma de U. brizantha antes do "
    "desenho final dos primers. Essa etapa não está implementada no pipeline atual — "
    "não é uma lacuna secundária, é um requisito do projeto ainda não cumprido.",
)
para(
    doc,
    "Recomendação: antes de encomendar síntese de qualquer dsRNA, executar checagem "
    "de off-target por janela deslizante (19–21 nt, mimetizando siRNAs processados) "
    "da sequência candidata contra os transcriptomas/genomas disponíveis de "
    "Urochloa, usando ferramenta dedicada (ex. si-Fi, OFF-Spotter adaptado, ou "
    "BLASTN de janela curta com word_size pequeno e filtro de baixa complexidade "
    "desativado).",
    bold=True,
)

heading(doc, "5.4 Ambiguidade de parálogos locus_009/locus_015", level=2)
para(
    doc,
    "locus_009 e locus_015 são quase idênticos (comprimento de ramo próximo de "
    "zero, RBH de locus_015 só fecha com Arabidopsis). Se o dsRNA for desenhado numa "
    "região conservada entre eles, pode silenciar mais de uma cópia sem que isso "
    "seja uma decisão intencional.",
)
para(
    doc,
    "Recomendação: alinhar as sequências nucleotídicas de locus_001, locus_009 e "
    "locus_015, marcar as posições diagnósticas (SNPs/indels) que distinguem "
    "locus_001 das demais, e decidir conscientemente entre silenciamento específico "
    "(região única) ou co-silenciamento intencional (região compartilhada) — "
    "documentando a escolha.",
    bold=True,
)

heading(doc, "5.5 Extensão exata do transcrito — PARCIALMENTE RESOLVIDO em 11/07/2026", level=2)
para(
    doc,
    "O modelo gênico original de locus_001 (4 éxons, 3 íntrons, stop códon no "
    "lugar esperado) veio inteiramente de alinhamento spliced-aware via miniprot, "
    "sem UTR modelada.",
)
para(
    doc,
    "Atualização (Seção 3.7): o gene homólogo real no subgenoma brizantha do "
    "assembly U. decumbens (URODEC1_LOCUS91967) tem anotação de mRNA completo com "
    "UTRs (1.222 nt totais, 525 nt de CDS, ~697 nt de UTR 5'+3' reais). A "
    "estrutura de 4 éxons/3 íntrons foi confirmada de forma independente.",
    bold=True,
)
para(
    doc,
    "O que ainda falta: não está confirmado se a anotação de UTR veio de "
    "evidência transcricional direta (RNA-seq) da própria U. decumbens/brizantha "
    "ou de modelo computacional. Recomendação atualizada: usar "
    "brizantha_mRNA_with_UTR.fasta (Seção 3.7) como ponto de partida para desenho "
    "de primers/dsRNA na região de UTR 3' (mais específica entre parálogos), mas "
    "confirmar por 5'/3' RACE ou RT-PCR ancorado em RNA real de U. brizantha cv. "
    "Marandu antes do desenho final.",
    bold=True,
)

heading(doc, "5.6 Gene SBP5 — bloqueio total do segundo alvo da proposta", level=2)
para(
    doc,
    "Nenhuma sequência de referência foi identificada ainda para o gene SBP5, o "
    "segundo alvo da proposta científica (junto com FT-like). Toda a etapa "
    "equivalente de descoberta de loci para esse gene está bloqueada até que uma "
    "sequência de referência seja definida.",
)

heading(doc, "5.7 Gargalos técnicos do pipeline", level=2)
bullet(
    doc,
    "Drift entre configuração e arquivo de entrada (encontrado e corrigido durante "
    "a auditoria que precedeu este relatório): o arquivo CDS de referência havia "
    "sido renomeado localmente sem atualizar a configuração do Nextflow, o que "
    "teria feito uma execução limpa falhar. Corrigido e commitado.",
)
bullet(
    doc,
    "Ausência de um critério de qualidade formal e ponderado antes de promover um "
    "locus a \"candidato confirmado\" — o índice usado no pipeline é uma soma "
    "não-ponderada de três indicadores binários. A confiança em locus_001 resultou "
    "de inspeção manual dos dados brutos, não de um critério automatizado.",
)
bullet(
    doc,
    "Dependência de um arquivo de referência (Pfam-A.hmm) compartilhado com outro "
    "projeto no mesmo servidor — fonte de possível quebra silenciosa se movido ou "
    "atualizado externamente.",
)

heading(doc, "5.8 Síntese das limitações", level=2)
para(
    doc,
    "locus_001 é um candidato forte e bem fundamentado dentro do que a "
    "bioinformática pode provar, e ficou mais forte ainda em 11/07/2026 com a "
    "confirmação em sequência real de brizantha (Seção 3.7): divergência de "
    "apenas 1 aminoácido e UTRs reais disponíveis, reduzindo dois gargalos "
    "(5.2, 5.5) de \"bloqueante\" para \"parcialmente resolvido\". Isso ainda não "
    "é o mesmo que \"pronto para desenho final de primers/dsRNA em bancada\". "
    "Seguem bloqueantes: 5.1 (validação experimental — nenhum PCR/RT-PCR real "
    "feito ainda) e 5.3 (off-target — ainda não avaliado, exigido pela "
    "proposta). Recomendação de sequência: PCR e Sanger em DNA genômico real da "
    "cv. Marandu usando a sequência de brizantha_CDS.fasta (Seção 3.7) como "
    "molde primeiro, depois off-target, antes de qualquer decisão de síntese de "
    "dsRNA.",
    bold=True,
)

# ---- 6. Conclusão ---------------------------------------------------------
heading(doc, "6. Conclusão", level=1)
para(
    doc,
    "A análise de bioinformática identificou e confirmou, com três linhas "
    "independentes de evidência (sequência, ortologia recíproca e filogenia), "
    "locus_001 como candidato a ortólogo do gene FLOWERING LOCUS T em Urochloa "
    "ruziziensis, adequado para orientar a próxima etapa de seleção de regiões "
    "conservadas para desenho de dsRNA (Objetivo Específico 1 da proposta "
    "científica). Um segundo achado — a reclassificação de locus_012 como "
    "ortólogo de MFT, não de FT — refina a compreensão da família gênica FT-like em "
    "Urochloa e evita um possível erro de direcionamento do dsRNA. Um terceiro "
    "achado, obtido no mesmo dia da revisão que originou este relatório, elevou "
    "a confiança no resultado de forma direta: a busca da sequência de locus_001 "
    "no subgenoma real de U. brizantha, presente no assembly publicado de "
    "U. decumbens cv. Basilisk, encontrou uma cópia 99,4% idêntica — divergência "
    "mínima e agora quantificada, com UTRs reais disponíveis pela primeira vez.",
)
para(
    doc,
    "Ao mesmo tempo, este relatório documenta de forma explícita e honesta os "
    "limites do que a bioinformática pode garantir sozinha: o resultado ainda é "
    "in silico e carece da análise de especificidade off-target exigida pela "
    "própria proposta, além de confirmação experimental direta na cultivar de "
    "trabalho (Marandu). A Seção 5 define um caminho concreto e sequencial para "
    "transformar este candidato confirmado por bioinformática — agora apoiado em "
    "sequência real de brizantha — em uma sequência segura para síntese de "
    "dsRNA em bancada.",
)

# ---- Referências ------------------------------------------------------------
doc.add_page_break()
heading(doc, "Referências", level=1)

refs = [
    "Allen, J. E., Pertea, M., & Salzberg, S. L. (2004). Computational Gene Prediction Using Multiple Sources of Evidence. Genome Research, 14(1), 142–148. https://doi.org/10.1101/gr.1562804",
    "Capella-Gutiérrez, S., Silla-Martínez, J. M., & Gabaldón, T. (2009). trimAl: a tool for automated alignment trimming in large-scale phylogenetic analyses. Bioinformatics, 25(15), 1972–1973. https://doi.org/10.1093/bioinformatics/btp348",
    "Dalquen, D. A., & Dessimoz, C. (2013). Bidirectional Best Hits Miss Many Orthologs in Duplication-Rich Clades such as Plants and Animals. Genome Biology and Evolution, 5(10), 1800–1806. https://doi.org/10.1093/gbe/evt132",
    "Danilevskaya, O. N., Meng, X., Hou, Z., Ananiev, E. V., & Simmons, C. R. (2008). A Genomic and Expression Compendium of the Expanded PEBP Gene Family from Maize. Plant Physiology, 146, 250–264. https://doi.org/10.1104/pp.107.109538",
    "de Jesus, D. A., Batista, D. M., Monteiro, E. F., Salzman, S., Carvalho, L. M., Santana, K., & André, T. (2022). Structural changes and adaptive evolutionary constraints in FLOWERING LOCUS T and TERMINAL FLOWER1-like genes of flowering plants. Frontiers in Genetics, 13, 954015. https://doi.org/10.3389/fgene.2022.954015",
    "Eddy, S. R. (2011). Accelerated Profile HMM Searches. PLOS Computational Biology, 7(10), e1002195. https://doi.org/10.1371/journal.pcbi.1002195",
    "Ferreira, R. C. U., et al. (2021). An Overview of the Genetics and Genomics of the Urochloa Species Most Commonly Used in Pastures. Frontiers in Plant Science, 12, 770461. https://doi.org/10.3389/fpls.2021.770461",
    "Hoang, D. T., Chernomor, O., von Haeseler, A., Minh, B. Q., & Vinh, L. S. (2018). UFBoot2: Improving the Ultrafast Bootstrap Approximation. Molecular Biology and Evolution, 35(2), 518–522. https://doi.org/10.1093/molbev/msx281",
    "Kalyaanamoorthy, S., Minh, B. Q., Wong, T. K. F., von Haeseler, A., & Jermiin, L. S. (2017). ModelFinder: fast model selection for accurate phylogenetic estimates. Nature Methods, 14(6), 587–589. https://doi.org/10.1038/nmeth.4285",
    "Karlgren, A., Gyllenstrand, N., Källman, T., Sundström, J. F., Moore, D., Lascoux, M., & Lagercrantz, U. (2011). Evolution of the PEBP Gene Family in Plants: Functional Diversification in Seed Plant Evolution. Plant Physiology, 156(4), 1967–1977. https://doi.org/10.1104/pp.111.176206",
    "Katoh, K., & Standley, D. M. (2013). MAFFT Multiple Sequence Alignment Software Version 7: Improvements in Performance and Usability. Molecular Biology and Evolution, 30(4), 772–780. https://doi.org/10.1093/molbev/mst010",
    "Kojima, S., Takahashi, Y., Kobayashi, Y., et al. (2002). Hd3a, a Rice Ortholog of the Arabidopsis FT Gene, Promotes Transition to Flowering Downstream of Hd1 under Short-Day Conditions. Plant and Cell Physiology, 43, 1096–1105. https://doi.org/10.1093/pcp/pcf156",
    "Li, H. (2023). Protein-to-genome alignment with miniprot. Bioinformatics, 39(1), btad014. https://doi.org/10.1093/bioinformatics/btad014",
    "Marcussen, T., Sandve, S. R., Heier, L., Spannagl, M., Pfeifer, M., The International Wheat Genome Sequencing Consortium, et al. (2014). Ancient hybridizations among the ancestral genomes of bread wheat. Science, 345(6194), 1250092. https://doi.org/10.1126/science.1250092",
    "Minh, B. Q., Schmidt, H. A., Chernomor, O., Schrempf, D., Woodhams, M. D., von Haeseler, A., & Lanfear, R. (2020). IQ-TREE 2: New Models and Efficient Methods for Phylogenetic Inference in the Genomic Era. Molecular Biology and Evolution, 37(5), 1530–1534. https://doi.org/10.1093/molbev/msaa015",
    "Mistry, J., Chuguransky, S., Williams, L., Qureshi, M., Salazar, G. A., Sonnhammer, E. L. L., et al. (2021). Pfam: The Protein Families Database in 2021. Nucleic Acids Research, 49(D1), D412–D419. https://doi.org/10.1093/nar/gkaa913",
    "Moreno-Hagelsieb, G., & Latimer, K. (2008). Choosing BLAST options for better detection of orthologs as reciprocal best hits. Bioinformatics, 24(3), 319–324. https://doi.org/10.1093/bioinformatics/btm585",
    "Tatusov, R. L., Koonin, E. V., & Lipman, D. J. (1997). A Genomic Perspective on Protein Families. Science, 278(5338), 631–637. https://doi.org/10.1126/science.278.5338.631",
    "Wickland, D. P., & Hanzawa, Y. (2015). The FLOWERING LOCUS T/TERMINAL FLOWER 1 Gene Family: Functional Evolution and Molecular Mechanisms. Molecular Plant, 8(7), 983–997. https://doi.org/10.1016/j.molp.2015.01.007",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(ref)
    run.font.size = Pt(10)

para(doc, "")
p = doc.add_paragraph()
run = p.add_run(
    "Nota metodológica sobre as referências: todas as 19 citações acima foram "
    "verificadas individualmente na fonte real (busca + confirmação de título, "
    "autores, ano e periódico) antes de entrar neste relatório ou no documento de "
    "trabalho artigo.md que o originou — nenhuma foi usada de memória sem "
    "confirmação. Itens metodológicos sem citação direta encontrada na literatura "
    "(ex. thresholds exatos de E-value) foram marcados no documento de trabalho como "
    "suporte parcial/de boas práticas, em vez de citação forçada."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = GREY

doc.save(OUT)
print(f"Salvo: {OUT}")
